"""
文本提取器模块

提供从各种文档格式中提取文本的功能。
支持的格式：PDF、Word (docx)、Excel (xlsx)
"""

from typing import Optional, Dict, Any
import logging
from abc import ABC, abstractmethod

logger = logging.getLogger(__name__)


class TextExtractor(ABC):
    """
    文本提取器基类
    
    定义文本提取的统一接口。
    """
    
    def __init__(self):
        """初始化文本提取器"""
        self.logger = logging.getLogger(self.__class__.__name__)
    
    @abstractmethod
    def extract(self, file_data: bytes, file_name: str) -> str:
        """
        从文件中提取文本
        
        Args:
            file_data: 文件的二进制数据
            file_name: 文件名（用于日志）
            
        Returns:
            str: 提取的文本内容
            
        Raises:
            Exception: 如果提取失败
        """
        pass
    
    @abstractmethod
    def can_extract(self, mime_type: str, extension: str) -> bool:
        """
        判断是否可以提取指定类型的文件
        
        Args:
            mime_type: 文件的 MIME 类型
            extension: 文件扩展名
            
        Returns:
            bool: 是否可以提取
        """
        pass


class PDFExtractor(TextExtractor):
    """
    PDF 文本提取器
    
    使用 pdfplumber 提取 PDF 文件中的文本内容。
    """
    
    SUPPORTED_MIME_TYPES = {'application/pdf'}
    SUPPORTED_EXTENSIONS = {'pdf'}
    
    def can_extract(self, mime_type: str, extension: str) -> bool:
        """判断是否可以提取 PDF 文件"""
        return (mime_type.lower() in self.SUPPORTED_MIME_TYPES or 
                extension.lower() in self.SUPPORTED_EXTENSIONS)
    
    def extract(self, file_data: bytes, file_name: str) -> str:
        """
        从 PDF 文件中提取文本
        
        Args:
            file_data: PDF 文件的二进制数据
            file_name: 文件名
            
        Returns:
            str: 提取的文本内容
            
        Raises:
            Exception: 如果提取失败
        """
        try:
            import pdfplumber
            import io
            
            # 使用 BytesIO 包装文件数据
            pdf_file = io.BytesIO(file_data)
            
            # 打开 PDF 文件
            with pdfplumber.open(pdf_file) as pdf:
                # 提取所有页面的文本
                text_parts = []
                for page_num, page in enumerate(pdf.pages, start=1):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                        else:
                            self.logger.debug(f"PDF 第 {page_num} 页无文本内容: {file_name}")
                    except Exception as e:
                        self.logger.warning(f"PDF 第 {page_num} 页提取失败: {e}")
                        continue
                
                # 合并所有页面的文本
                full_text = '\n\n'.join(text_parts)
                
                if not full_text.strip():
                    raise ValueError("PDF 文件中没有可提取的文本内容（可能是扫描件或图片）")
                
                self.logger.info(f"PDF 文本提取成功: {file_name} ({len(full_text)} 字符, {len(pdf.pages)} 页)")
                return full_text
                
        except ImportError:
            self.logger.error("pdfplumber 未安装，无法提取 PDF 文本")
            raise Exception("PDF 文本提取功能不可用：缺少 pdfplumber 库")
        except Exception as e:
            self.logger.error(f"PDF 文本提取失败: {e}", exc_info=True)
            raise


class DocxExtractor(TextExtractor):
    """
    Word 文档文本提取器
    
    使用 python-docx 提取 Word 文档中的文本内容。
    """
    
    SUPPORTED_MIME_TYPES = {
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'application/msword'
    }
    SUPPORTED_EXTENSIONS = {'docx', 'doc'}
    
    def can_extract(self, mime_type: str, extension: str) -> bool:
        """判断是否可以提取 Word 文档"""
        return (mime_type.lower() in self.SUPPORTED_MIME_TYPES or 
                extension.lower() in self.SUPPORTED_EXTENSIONS)
    
    def extract(self, file_data: bytes, file_name: str) -> str:
        """
        从 Word 文档中提取文本
        
        Args:
            file_data: Word 文档的二进制数据
            file_name: 文件名
            
        Returns:
            str: 提取的文本内容
            
        Raises:
            Exception: 如果提取失败
        """
        try:
            from docx import Document
            import io
            
            # 使用 BytesIO 包装文件数据
            docx_file = io.BytesIO(file_data)
            
            # 打开 Word 文档
            doc = Document(docx_file)
            
            # 提取所有段落的文本
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # 提取表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            # 合并所有文本
            full_text = '\n'.join(text_parts)
            
            if not full_text.strip():
                raise ValueError("Word 文档中没有可提取的文本内容")
            
            self.logger.info(f"Word 文档文本提取成功: {file_name} ({len(full_text)} 字符, {len(doc.paragraphs)} 段落)")
            return full_text
            
        except ImportError:
            self.logger.error("python-docx 未安装，无法提取 Word 文档文本")
            raise Exception("Word 文档文本提取功能不可用：缺少 python-docx 库")
        except Exception as e:
            self.logger.error(f"Word 文档文本提取失败: {e}", exc_info=True)
            raise


class XlsxExtractor(TextExtractor):
    """
    Excel 文档文本提取器
    
    使用 openpyxl 提取 Excel 文档中的文本内容。
    """
    
    SUPPORTED_MIME_TYPES = {
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'application/vnd.ms-excel'
    }
    SUPPORTED_EXTENSIONS = {'xlsx', 'xls'}
    
    def can_extract(self, mime_type: str, extension: str) -> bool:
        """判断是否可以提取 Excel 文档"""
        return (mime_type.lower() in self.SUPPORTED_MIME_TYPES or 
                extension.lower() in self.SUPPORTED_EXTENSIONS)
    
    def extract(self, file_data: bytes, file_name: str) -> str:
        """
        从 Excel 文档中提取文本
        
        Args:
            file_data: Excel 文档的二进制数据
            file_name: 文件名
            
        Returns:
            str: 提取的文本内容
            
        Raises:
            Exception: 如果提取失败
        """
        try:
            from openpyxl import load_workbook
            import io
            
            # 使用 BytesIO 包装文件数据
            xlsx_file = io.BytesIO(file_data)
            
            # 打开 Excel 文档
            wb = load_workbook(xlsx_file, data_only=True)
            
            # 提取所有工作表的文本
            text_parts = []
            total_rows = 0
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text_parts.append(f"=== 工作表: {sheet_name} ===")
                
                # 提取每一行的数据
                for row in sheet.iter_rows(values_only=True):
                    # 过滤空值并转换为字符串
                    row_values = [str(cell) if cell is not None else '' for cell in row]
                    # 过滤空行
                    if any(val.strip() for val in row_values):
                        text_parts.append(' | '.join(row_values))
                        total_rows += 1
            
            # 合并所有文本
            full_text = '\n'.join(text_parts)
            
            if not full_text.strip() or total_rows == 0:
                raise ValueError("Excel 文档中没有可提取的文本内容")
            
            self.logger.info(f"Excel 文档文本提取成功: {file_name} ({len(full_text)} 字符, {total_rows} 行, {len(wb.sheetnames)} 个工作表)")
            return full_text
            
        except ImportError:
            self.logger.error("openpyxl 未安装，无法提取 Excel 文档文本")
            raise Exception("Excel 文档文本提取功能不可用：缺少 openpyxl 库")
        except Exception as e:
            self.logger.error(f"Excel 文档文本提取失败: {e}", exc_info=True)
            raise


class TextExtractorFactory:
    """
    文本提取器工厂
    
    根据文件类型选择合适的文本提取器。
    """
    
    def __init__(self):
        """初始化文本提取器工厂"""
        self.extractors = [
            PDFExtractor(),
            DocxExtractor(),
            XlsxExtractor()
        ]
        self.logger = logging.getLogger(self.__class__.__name__)
    
    def get_extractor(self, mime_type: str, extension: str) -> Optional[TextExtractor]:
        """
        获取适合指定文件类型的文本提取器
        
        Args:
            mime_type: 文件的 MIME 类型
            extension: 文件扩展名
            
        Returns:
            Optional[TextExtractor]: 文本提取器，如果没有合适的提取器则返回 None
        """
        for extractor in self.extractors:
            if extractor.can_extract(mime_type, extension):
                self.logger.debug(f"选择文本提取器: {extractor.__class__.__name__} (mime_type={mime_type}, extension={extension})")
                return extractor
        
        self.logger.debug(f"没有找到合适的文本提取器: mime_type={mime_type}, extension={extension}")
        return None
    
    def extract_text(self, file_data: bytes, file_name: str, mime_type: str, extension: str) -> Optional[str]:
        """
        提取文件中的文本
        
        Args:
            file_data: 文件的二进制数据
            file_name: 文件名
            mime_type: 文件的 MIME 类型
            extension: 文件扩展名
            
        Returns:
            Optional[str]: 提取的文本内容，如果提取失败则返回 None
        """
        extractor = self.get_extractor(mime_type, extension)
        if extractor is None:
            self.logger.warning(f"没有找到合适的文本提取器: {file_name}")
            return None
        
        try:
            return extractor.extract(file_data, file_name)
        except Exception as e:
            self.logger.error(f"文本提取失败: {file_name}, 错误: {e}")
            return None
